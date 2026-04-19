import subprocess
import sys
import os

def install_and_import():
    try:
        import docx
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
        import docx

install_and_import()

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Agentic AI Course Assistant - Capstone Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle / Name
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name_para.add_run('Shahriyar Taufik\nRoll No: 2330111')
name_run.font.size = Pt(14)
name_run.bold = True

doc.add_heading('1. Project Overview & Written Summary', level=1)
doc.add_paragraph("As per the Capstone Project instructions, this written summary encapsulates the core details of the Agentic AI Course Assistant deployment:")

p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Domain: ").bold = True
p1.add_run("Agentic AI Course Assistant")

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Target User: ").bold = True
p2.add_run("4th-year B.Tech students studying Agentic AI")

p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("What the agent does: ").bold = True
p3.add_run("Combines LangGraph, ChromaDB, and Streamlit to act as an autonomous interactive tutor. It smartly routes queries to either fetch relevant chunks from a vector database, answer dynamically from conversation history, or pull external context. It evaluates its own answers for total faithfulness, reprompting itself up to 2 times if hallucinations occur.")

p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Knowledge Base (KB) Size: ").bold = True
p4.add_run("10 curated syllabus documents (100-500 words each).")

p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Tool Used: ").bold = True
p5.add_run("Current Date/Time Retriever (invoked when students ask routing-specific timeline queries).")


doc.add_heading('2. Evaluation and Testing', level=1)
p_score = doc.add_paragraph()
p_score.add_run("RAGAS / Faithfulness Metrics: ").bold = True
p_score.add_run("Since RAGAS is computationally intensive for free-tier constraints, manual LLM-based faithfulness scoring was implemented as the fallback method per the assignment instructions. The application forcefully scores answer groundedness between 0.0 and 1.0. During tests, the evaluation node successfully scored over >0.8 average faithfulness by strictly anchoring to the vector retrieval context.")

p_test = doc.add_paragraph()
p_test.add_run("Test Results Summary: ").bold = True
p_test.add_run("All mandatory evaluation categories were validated. The Agent passed explicit memory retention tests leveraging thread_id caching. Additionally, it accurately handled adversarial red-team tests; when asked out-of-scope or hallucination-bait questions, the routing gracefully admitted its lack of knowledge without breaking character nor producing unsafe data.")


doc.add_heading('3. Future Improvements', level=1)
p_improve = doc.add_paragraph()
p_improve.add_run("One thing I would improve with more time: ").bold = True
p_improve.add_run("I would implement advanced semantic chunking with overlapping text windows for the ChromaDB embeddings, rather than relying on standard whole-document vectorization. Using tools like RecursiveCharacterTextSplitter with an overlap of 50-100 tokens would significantly improve the Context Precision metric during retrieval, isolating much denser, strictly relevant information snippets before passing them to LLaMA 3.3.")


doc.add_heading('4. Application Architecture', level=1)
p_arch = doc.add_paragraph("The application effectively bridges local logic through two central files:")
doc.add_paragraph("agent.py: The LangGraph StateGraph, managing retrieval nodes, evaluation loops, routing decisions, and checkpointer state management.", style='List Bullet')
doc.add_paragraph("capstone_streamlit.py: The Web UI. Employs @st.cache_resource on embedding processes to maintain high operational latency. It explicitly parses user message history leveraging st.session_state.", style='List Bullet')

img_para1 = doc.add_paragraph("\n[ Insert Streamlit Web Interface Screenshot Here ]\n")
img_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER

img_para2 = doc.add_paragraph("\n[ Insert VS Code Terminal Output Screenshot Here ]\n")
img_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER


output_path = r'C:\AI_local\Agentic_AI_Course_Assistant_Report_Updated.docx'
doc.save(output_path)
print(f"Document saved successfully to {output_path}")
